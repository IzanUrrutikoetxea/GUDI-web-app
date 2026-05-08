"""
Generates informe.docx for the GUDI project.
Run: python docs/generate_informe.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

FONT = "Times New Roman"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, size=12,
                  bold=False, italic=False, space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 14, 3: 12}
    p = add_paragraph(doc, text, size=sizes.get(level, 12), bold=True,
                      space_before=14, space_after=6)
    return p

def add_bullet(doc, text, size=12):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_page_break(doc):
    doc.add_page_break()

def add_page_numbers(section):
    """Add page number bottom-right to footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = p.add_run()
    run.font.name = FONT
    run.font.size = Pt(10)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def set_margins(section, top=2.5, bottom=2.5, left=3, right=2.5):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)

# ─── Document ─────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
set_margins(section)
add_page_numbers(section)

# ══════════════════════════════════════════════════════════════════════════════
# 1. PORTADA
# ══════════════════════════════════════════════════════════════════════════════

for _ in range(7):
    add_paragraph(doc, "")

add_paragraph(doc, "GUDI", align=WD_ALIGN_PARAGRAPH.CENTER, size=36, bold=True,
              color=(37, 99, 235))
add_paragraph(doc, "Gestión Unificada Digital Inteligente",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=16, italic=True, space_after=30)

add_paragraph(doc, "Informe técnico del proyecto",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=40)

add_paragraph(doc, "Asignatura: Sistemas Web",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=6)
add_paragraph(doc, "Autores:",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, space_after=4)
add_paragraph(doc, "Izan Urrutikoetxea",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=4)
add_paragraph(doc, "Gorka Delgado",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=4)
add_paragraph(doc, "Mario Romero",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=20)
add_paragraph(doc, "Profesor: Javier del Valle Echevarri",
              align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=6)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════

add_paragraph(doc, "Índice", align=WD_ALIGN_PARAGRAPH.CENTER,
              size=18, bold=True, space_after=20)

toc = [
    ("1.", "Introducción", "3"),
    ("2.", "Arquitectura general del sistema", "4"),
    ("3.", "Stack técnico y justificación", "5"),
    ("4.", "Estructura del repositorio", "7"),
    ("5.", "Decisiones técnicas", "8"),
    ("  5.1.", "Frontend", "8"),
    ("  5.2.", "Backend", "9"),
    ("  5.3.", "Base de datos", "10"),
    ("  5.4.", "Autenticación y seguridad", "10"),
    ("  5.5.", "Contenerización con Docker", "11"),
    ("6.", "Módulos funcionales", "12"),
    ("7.", "Observabilidad", "14"),
    ("  7.1.", "Métricas — Prometheus y Grafana", "14"),
    ("  7.2.", "Trazas distribuidas — OpenTelemetry y Jaeger", "15"),
    ("  7.3.", "Logs estructurados — Stack ELK", "15"),
    ("8.", "Comunicación entre servicios — gRPC", "16"),
    ("9.", "Flujo de trabajo con Git y GitHub", "17"),
    ("10.", "Conclusiones", "18"),
]

for num, title, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(13), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = p.add_run(f"{num}  {title}")
    set_font(r1, size=12)
    r2 = p.add_run(f"\t{page}")
    set_font(r2, size=12)

add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. CONTENIDO
# ══════════════════════════════════════════════════════════════════════════════

# ── 1. Introducción ───────────────────────────────────────────────────────────
add_heading(doc, "1. Introducción")

add_paragraph(doc, (
    "El presente informe documenta el desarrollo del proyecto GUDI (Gestión Unificada Digital "
    "Inteligente), una aplicación web diseñada para la digitalización y automatización de "
    "procesos en pequeñas y medianas empresas (PYMEs). El proyecto ha sido desarrollado en el "
    "marco de la asignatura Sistemas Web como proyecto de curso por un equipo de tres "
    "desarrolladores."
), space_after=8)

add_paragraph(doc, (
    "El objetivo principal de GUDI es ofrecer a las PYMEs una plataforma unificada desde la que "
    "gestionar sus citas, presupuestos y comunicaciones con clientes, todo ello accesible desde "
    "un navegador web sin necesidad de instalar software adicional."
), space_after=8)

add_paragraph(doc, "Los objetivos técnicos del proyecto son:", space_after=4)
add_bullet(doc, "Diseñar e implementar una arquitectura cliente-servidor moderna y escalable.")
add_bullet(doc, "Aplicar buenas prácticas de desarrollo: separación de responsabilidades, control de versiones y documentación.")
add_bullet(doc, "Garantizar la seguridad mediante autenticación basada en JWT.")
add_bullet(doc, "Contenerizar la aplicación completa con Docker para facilitar el despliegue.")
add_bullet(doc, "Incorporar un sistema de observabilidad completo: métricas, trazas y logs.")
add_bullet(doc, "Experimentar con arquitecturas de microservicios mediante comunicación gRPC.")
add_paragraph(doc, "", space_after=4)

add_page_break(doc)

# ── 2. Arquitectura general ───────────────────────────────────────────────────
add_heading(doc, "2. Arquitectura general del sistema")

add_paragraph(doc, (
    "GUDI sigue una arquitectura de tres capas clásica en aplicaciones web modernas, complementada "
    "con un microservicio independiente para la gestión de notificaciones y un stack de "
    "observabilidad completo."
), space_after=8)

add_heading(doc, "Capas principales", level=2)
add_bullet(doc, "Capa de presentación (Frontend): aplicación React ejecutada en el navegador del usuario, que se comunica con el backend a través de una API REST.")
add_bullet(doc, "Capa de aplicación (Backend): servidor Node.js con Express que expone la API REST, aplica la lógica de negocio y actúa como API Gateway hacia el microservicio gRPC.")
add_bullet(doc, "Capa de persistencia (Base de datos): PostgreSQL como sistema de gestión de base de datos relacional, accedido a través del ORM Prisma.")
add_paragraph(doc, "", space_after=4)

add_heading(doc, "Microservicio de notificaciones", level=2)
add_paragraph(doc, (
    "El sistema incorpora un microservicio independiente (notifications-service) que se comunica "
    "con el backend mediante el protocolo gRPC. Este servicio gestiona el envío y las estadísticas "
    "de notificaciones, permitiendo desacoplar esta responsabilidad del backend principal."
), space_after=8)

add_heading(doc, "Stack de observabilidad", level=2)
add_paragraph(doc, (
    "Para monitorizar el sistema en tiempo real se han integrado tres pilares de observabilidad: "
    "métricas con Prometheus y Grafana, trazas distribuidas con OpenTelemetry y Jaeger, y logs "
    "estructurados con el stack ELK (Elasticsearch, Logstash, Kibana)."
), space_after=8)

add_paragraph(doc, (
    "Todos los componentes se despliegan mediante Docker Compose, lo que permite levantar el "
    "entorno completo con un único comando, independientemente del sistema operativo del "
    "desarrollador o del servidor de despliegue."
), space_after=8)

add_page_break(doc)

# ── 3. Stack técnico ─────────────────────────────────────────────────────────
add_heading(doc, "3. Stack técnico y justificación")

add_paragraph(doc, (
    "La elección del stack tecnológico ha sido una de las decisiones más importantes del proyecto. "
    "A continuación se detalla cada tecnología seleccionada y la razón que motivó su elección."
), space_after=10)

stack = [
    ("React 18 + TypeScript (Frontend)",
     "React es la biblioteca de construcción de interfaces de usuario más utilizada en la industria. "
     "Su modelo basado en componentes reutilizables facilita la organización del código y permite "
     "desarrollar interfaces dinámicas de forma eficiente. Se eligió frente a Vue o Angular por la "
     "mayor familiaridad del equipo y su ecosistema más maduro. TypeScript se incorporó para añadir "
     "tipado estático, lo que reduce errores en tiempo de desarrollo y mejora la mantenibilidad, "
     "especialmente en proyectos con múltiples desarrolladores."),
    ("Vite (Herramienta de construcción)",
     "Vite fue seleccionado como bundler y servidor de desarrollo frente a la alternativa clásica "
     "Create React App (CRA). Vite ofrece tiempos de arranque del servidor de desarrollo "
     "drásticamente menores gracias al uso de ES Modules nativos, y builds de producción más "
     "rápidos con Rollup. En un proyecto educativo donde el equipo arranca y para el servidor "
     "frecuentemente, esta mejora de experiencia de desarrollo es especialmente valorable."),
    ("Node.js + Express + TypeScript (Backend)",
     "Node.js permite utilizar JavaScript en el servidor, lo que facilita compartir conocimiento "
     "y tipos entre frontend y backend. Express es el framework web más extendido para Node.js, "
     "con una curva de aprendizaje mínima y una arquitectura minimalista que no impone restricciones "
     "excesivas. Se valoraron alternativas como NestJS (más estructurado) y Fastify (más rápido), "
     "pero Express fue elegido por su simplicidad y porque el equipo puede implementar la "
     "arquitectura deseada sin las imposiciones de un framework más opinionado."),
    ("PostgreSQL (Base de datos)",
     "PostgreSQL es el sistema de gestión de bases de datos relacional de código abierto más "
     "robusto y completo disponible. Se eligió frente a MySQL por su mejor soporte de "
     "características avanzadas de SQL, mejor manejo de tipos JSON, y su reputación de "
     "fiabilidad en entornos de producción. Frente a alternativas NoSQL como MongoDB, se optó "
     "por una base de datos relacional porque los datos del sistema (usuarios, citas, "
     "presupuestos) tienen relaciones claras y definidas que se benefician del modelo relacional."),
    ("Prisma (ORM)",
     "Prisma es un ORM moderno para Node.js y TypeScript que genera un cliente con tipado "
     "completo a partir del esquema de la base de datos. Esto permite detectar errores de "
     "consulta en tiempo de compilación y proporciona autocompletado en el editor. Se eligió "
     "frente a Sequelize o TypeORM por su mejor integración con TypeScript, su sistema de "
     "migraciones más intuitivo y su Prisma Studio, que permite explorar la base de datos "
     "visualmente durante el desarrollo."),
    ("React Router v6 (Enrutamiento frontend)",
     "React Router es la solución de enrutamiento estándar en el ecosistema React. La versión 6 "
     "introduce una API más declarativa y soporte mejorado para layouts anidados, lo que se "
     "aprovecha en el proyecto para el layout principal con sidebar. Se implementaron rutas "
     "protegidas (PrivateRoute) y rutas públicas (PublicRoute) para gestionar el acceso según "
     "el estado de autenticación."),
    ("JWT — JSON Web Tokens (Autenticación)",
     "JWT fue elegido como mecanismo de autenticación por ser un estándar ampliamente adoptado "
     "para aplicaciones SPA (Single Page Applications). A diferencia de las sesiones del lado "
     "del servidor, los tokens JWT son stateless, lo que simplifica el escalado horizontal y "
     "elimina la necesidad de un almacén de sesiones compartido. Los tokens se almacenan en "
     "localStorage y se envían en el header Authorization de cada petición protegida."),
    ("Docker y Docker Compose (Contenerización)",
     "Docker permite empaquetar cada servicio con todas sus dependencias, garantizando que el "
     "entorno de desarrollo y el de producción sean idénticos. Docker Compose orquesta todos los "
     "servicios con un único fichero de configuración, lo que simplifica enormemente el proceso "
     "de onboarding para nuevos desarrolladores. Se utilizan perfiles de Compose para hacer el "
     "stack ELK opcional, ya que requiere más recursos de memoria."),
]

for title, desc in stack:
    add_paragraph(doc, title, size=12, bold=True, space_before=8, space_after=3)
    add_paragraph(doc, desc, space_after=6)

add_page_break(doc)

# ── 4. Estructura del repositorio ────────────────────────────────────────────
add_heading(doc, "4. Estructura del repositorio")

add_paragraph(doc, (
    "El repositorio sigue una estructura monorepo organizada por responsabilidades. Esta decisión "
    "facilita la visión global del proyecto y simplifica la gestión de dependencias compartidas, "
    "aunque cada capa puede desarrollarse e incluso desplegarse de forma independiente."
), space_after=8)

repo_structure = [
    ("frontend/", "Aplicación cliente React + TypeScript. Contiene páginas, componentes, servicios de API, contextos de estado, hooks personalizados y tipos TypeScript."),
    ("backend/", "Servidor Node.js + Express. Organizado en capas: routes (endpoints), controllers (manejo de petición/respuesta), services (lógica de negocio), middlewares (auth, métricas, logging) y config (Prisma, env, gRPC client)."),
    ("backend/prisma/", "Esquema de base de datos y migraciones gestionadas por Prisma. El fichero schema.prisma define todos los modelos: User, Appointment, Budget, BudgetItem y Message."),
    ("notifications-service/", "Microservicio gRPC independiente. Expone dos RPCs: Send (enviar notificación) y GetStats (estadísticas de envíos)."),
    ("infra/", "Ficheros de configuración de la infraestructura de observabilidad: prometheus.yml, otel-collector-config.yml, logstash.conf y datasources de Grafana."),
    ("docs/", "Documentación técnica del proyecto, incluyendo arquitectura, roadmap, normas de desarrollo, metodología y el presente informe."),
]

for folder, desc in repo_structure:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(folder + "  ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

add_paragraph(doc, "", space_after=4)

add_paragraph(doc, (
    "La decisión de mantener frontend, backend y microservicio en el mismo repositorio responde "
    "a criterios pedagógicos: facilita la revisión del código por el profesor y el seguimiento "
    "del progreso del equipo, y simplifica el proceso de CI/CD en el contexto de un proyecto "
    "educativo. En un entorno de producción real con equipos grandes, cada servicio podría "
    "tener su propio repositorio."
), space_after=8)

add_page_break(doc)

# ── 5. Decisiones técnicas ────────────────────────────────────────────────────
add_heading(doc, "5. Decisiones técnicas")

add_heading(doc, "5.1. Frontend", level=2)

add_paragraph(doc, "Gestión del estado de autenticación", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "Se optó por el Context API nativo de React en lugar de soluciones externas como Redux o "
    "Zustand. El estado de autenticación (token JWT y datos del usuario) se gestiona en un "
    "AuthContext que envuelve toda la aplicación. Esta decisión se justifica porque el estado "
    "compartido del proyecto es mínimo y no requiere la complejidad de un gestor de estado "
    "externo. El token persiste en localStorage para sobrevivir a recargas de página."
), space_after=8)

add_paragraph(doc, "Proxy en desarrollo vs. producción", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "El servidor de desarrollo de Vite incluye un proxy configurable que redirige las peticiones "
    "a /api hacia el backend. Esto resuelve el problema de CORS durante el desarrollo sin "
    "necesidad de configurar cabeceras adicionales en el backend. En el contexto Docker, la "
    "variable de entorno API_TARGET apunta al nombre del servicio backend (http://backend:3000), "
    "usando la red interna de Docker Compose en lugar de localhost."
), space_after=8)

add_paragraph(doc, "Rutas protegidas", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "Se implementaron componentes PrivateRoute y PublicRoute como wrappers de React Router. "
    "PrivateRoute redirige a /login si el usuario no está autenticado; PublicRoute redirige al "
    "dashboard si ya tiene sesión. Este patrón evita que usuarios no autenticados accedan a "
    "páginas privadas y que usuarios autenticados vean las pantallas de login/registro."
), space_after=8)

add_heading(doc, "5.2. Backend", level=2)

add_paragraph(doc, "Arquitectura en capas", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "El backend sigue el patrón Route → Controller → Service, que separa las responsabilidades "
    "en tres niveles. Las rutas definen los endpoints y aplican middlewares; los controllers "
    "reciben la petición, validan los parámetros y devuelven la respuesta; los services contienen "
    "la lógica de negocio y acceden a la base de datos a través de Prisma. Esta separación "
    "facilita el testeo unitario de la lógica de negocio de forma independiente al protocolo HTTP."
), space_after=8)

add_paragraph(doc, "Manejo centralizado de errores", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "Se implementó un middleware de manejo de errores (errorHandler) que captura cualquier "
    "excepción no controlada y devuelve una respuesta JSON consistente. Esto evita que el "
    "servidor devuelva respuestas HTML de error de Express por defecto y garantiza que el "
    "frontend siempre reciba errores en un formato predecible."
), space_after=8)

add_paragraph(doc, "Variables de entorno", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "Todas las configuraciones sensibles (cadena de conexión a la base de datos, clave JWT, "
    "URLs de servicios externos) se gestionan mediante variables de entorno, nunca hardcodeadas "
    "en el código. El fichero .env.example documenta las variables necesarias sin exponer "
    "valores reales. El fichero .env real está incluido en .gitignore para evitar su publicación "
    "accidental en el repositorio."
), space_after=8)

add_heading(doc, "5.3. Base de datos", level=2)

add_paragraph(doc, "Migraciones con Prisma", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "El esquema de base de datos evoluciona mediante migraciones gestionadas por Prisma. Cada "
    "cambio en el esquema genera automáticamente un fichero SQL versionado en el directorio "
    "prisma/migrations. En el entorno Docker, el comando npx prisma migrate deploy aplica las "
    "migraciones pendientes al arrancar el contenedor, garantizando que la base de datos siempre "
    "esté sincronizada con el código."
), space_after=8)

add_paragraph(doc, "Seeding de datos", size=12, bold=True, space_before=6, space_after=3)
add_paragraph(doc, (
    "Se creó un script de seed (prisma/seed.ts) que inicializa la base de datos con datos de "
    "prueba representativos: un usuario administrador, tres usuarios regulares, siete citas, "
    "cuatro presupuestos en diferentes estados y ocho mensajes de distintos canales. Esto "
    "facilita la demostración del sistema y las pruebas manuales."
), space_after=8)

add_heading(doc, "5.4. Autenticación y seguridad", level=2)

add_paragraph(doc, (
    "La autenticación se implementa mediante JSON Web Tokens (JWT). Al registrarse o iniciar "
    "sesión, el servidor devuelve un token firmado con el JWT_SECRET que incluye el ID y rol "
    "del usuario. Este token tiene una validez de 7 días. El middleware authMiddleware verifica "
    "la firma del token en cada petición protegida y rechaza con 401 las peticiones con tokens "
    "inválidos o expirados."
), space_after=8)

add_paragraph(doc, (
    "Las contraseñas se almacenan hasheadas con bcrypt usando un coste de 10 rondas, lo que "
    "hace computacionalmente inviable la obtención de la contraseña original incluso si la "
    "base de datos se viera comprometida. Nunca se devuelve el hash de la contraseña en "
    "ninguna respuesta de la API."
), space_after=8)

add_heading(doc, "5.5. Contenerización con Docker", level=2)

add_paragraph(doc, (
    "Cada servicio (backend, frontend, notifications-service) tiene su propio Dockerfile que "
    "describe cómo construir su imagen. El fichero docker-compose.yml orquesta todos los "
    "servicios, define las redes internas, los volúmenes de persistencia y las dependencias "
    "entre servicios (depends_on con healthcheck para garantizar que PostgreSQL esté listo "
    "antes de que el backend intente conectarse)."
), space_after=8)

add_paragraph(doc, (
    "El stack ELK (Elasticsearch, Logstash, Kibana) se declara con el perfil elk en "
    "docker-compose.yml, ya que requiere una cantidad significativa de memoria RAM. Esto permite "
    "que los desarrolladores con máquinas menos potentes puedan ejecutar el sistema sin el stack "
    "ELK activando únicamente los servicios esenciales."
), space_after=8)

add_page_break(doc)

# ── 6. Módulos funcionales ────────────────────────────────────────────────────
add_heading(doc, "6. Módulos funcionales")

add_paragraph(doc, (
    "El desarrollo del sistema se organizó en seis fases o módulos funcionales, cada uno "
    "desarrollado en una rama feature independiente y fusionado a develop mediante Pull Request."
), space_after=8)

modules = [
    ("Fase 1 — Preparación del entorno",
     "Configuración del repositorio GitHub, estructura de carpetas, inicialización de los proyectos "
     "Node.js y React con TypeScript, configuración de Prisma, creación de la primera migración y "
     "configuración de Docker. Se establecieron las bases técnicas sobre las que se construirían "
     "todos los módulos posteriores."),
    ("Fase 2 — Autenticación y usuarios",
     "Implementación del registro e inicio de sesión de usuarios. El backend expone los endpoints "
     "POST /api/auth/register y POST /api/auth/login, que devuelven un JWT. El frontend muestra "
     "páginas de login y registro con validación de formularios. Se implementó el sistema de roles "
     "(ADMIN / USER) mediante un enum en Prisma."),
    ("Fase 3 — Panel de control (Dashboard)",
     "El dashboard es la pantalla principal del sistema autenticado. Muestra estadísticas de usuarios "
     "(totales, administradores, nuevos este mes) obtenidas del endpoint GET /api/dashboard/stats. "
     "También incluye una lista de actividad reciente con los últimos usuarios registrados. Se implementó "
     "el layout general de la aplicación con sidebar de navegación."),
    ("Fase 4 — Módulo de agenda",
     "Permite gestionar citas y eventos. El backend implementa un CRUD completo sobre el modelo "
     "Appointment. El frontend ofrece una lista de citas con formulario modal para crear y editar, "
     "validación de fechas (fecha fin posterior a fecha inicio) y confirmación antes de eliminar."),
    ("Fase 5 — Presupuestación",
     "Sistema de generación y gestión de presupuestos para clientes. Los presupuestos contienen "
     "líneas de conceptos con cantidad y precio unitario, cálculo automático de base imponible, "
     "IVA y total en tiempo real. Los presupuestos tienen estados (Borrador, Enviado, Aceptado, "
     "Rechazado) y se puede generar una vista de documento imprimible."),
    ("Fase 6 — Bandeja omnicanal",
     "Centraliza las comunicaciones de la empresa en tres canales: Email, WhatsApp e Interno. "
     "La interfaz sigue el patrón de cliente de correo con lista de mensajes a la izquierda y "
     "vista de detalle a la derecha. Los mensajes tienen estados (No leído, Leído, Archivado) y "
     "se pueden filtrar por canal y estado. Los mensajes se marcan como leídos automáticamente "
     "al abrirlos."),
]

for title, desc in modules:
    add_paragraph(doc, title, size=12, bold=True, space_before=8, space_after=3)
    add_paragraph(doc, desc, space_after=6)

add_page_break(doc)

# ── 7. Observabilidad ─────────────────────────────────────────────────────────
add_heading(doc, "7. Observabilidad")

add_paragraph(doc, (
    "A petición del profesor, se implementó un sistema de observabilidad completo que cubre los "
    "tres pilares fundamentales del monitoring moderno: métricas, trazas y logs. Esta adición "
    "transforma el proyecto de una aplicación web simple a un sistema observable y monitorizable "
    "de forma profesional."
), space_after=8)

add_heading(doc, "7.1. Métricas — Prometheus y Grafana", level=2)

add_paragraph(doc, (
    "Se integró la librería prom-client en el backend para exponer métricas en formato Prometheus "
    "en el endpoint GET /metrics. Las métricas implementadas son:"
), space_after=4)

add_bullet(doc, "gudi_http_requests_total: contador de peticiones HTTP clasificadas por método, ruta y código de estado.")
add_bullet(doc, "gudi_http_request_duration_seconds: histograma de latencia de peticiones con buckets de 5ms a 2.5s.")
add_bullet(doc, "gudi_http_active_requests: gauge del número de peticiones activas en un momento dado.")
add_bullet(doc, "Métricas por defecto de Node.js: uso de CPU, memoria heap, event loop lag, etc.")
add_paragraph(doc, "", space_after=4)

add_paragraph(doc, (
    "Prometheus se configura para scrapear el endpoint /metrics del backend cada 15 segundos. "
    "Grafana se provisiona automáticamente con Prometheus como datasource mediante ficheros de "
    "configuración en infra/grafana/provisioning, evitando configuración manual al arrancar el "
    "sistema."
), space_after=8)

add_heading(doc, "7.2. Trazas distribuidas — OpenTelemetry y Jaeger", level=2)

add_paragraph(doc, (
    "Se instrumentó el backend con el SDK de OpenTelemetry utilizando auto-instrumentaciones "
    "que capturan automáticamente trazas de peticiones HTTP Express, consultas Prisma y llamadas "
    "gRPC. La inicialización del SDK se realiza en el fichero tracing.ts, que debe importarse "
    "como primera línea del punto de entrada del servidor para garantizar que todas las "
    "librerías queden instrumentadas desde el inicio."
), space_after=8)

add_paragraph(doc, (
    "Las trazas se exportan al OTel Collector mediante OTLP sobre gRPC, y el Collector las "
    "reenvía a Jaeger. La interfaz de Jaeger permite visualizar el recorrido completo de cada "
    "petición, incluyendo los tiempos de cada operación (consulta a base de datos, llamada al "
    "microservicio gRPC, etc.), lo que facilita la detección de cuellos de botella."
), space_after=8)

add_heading(doc, "7.3. Logs estructurados — Stack ELK", level=2)

add_paragraph(doc, (
    "Se integró Winston como sistema de logging estructurado. Todos los logs se emiten en "
    "formato JSON con campos estandarizados (timestamp, level, service, message) que pueden "
    "ser ingestados directamente por Logstash. El middleware requestLogger registra cada "
    "petición HTTP con su método, ruta, código de estado, duración y dirección IP del cliente."
), space_after=8)

add_paragraph(doc, (
    "El stack ELK se compone de Elasticsearch (almacenamiento), Logstash (ingesta y "
    "transformación) y Kibana (visualización). Logstash escucha en el puerto TCP 5000 "
    "en formato JSON Lines y crea índices diarios en Elasticsearch (gudi-logs-YYYY.MM.DD). "
    "Dado el alto consumo de memoria de Elasticsearch (mínimo 1 GB), el stack ELK se declara "
    "como perfil opcional en docker-compose.yml."
), space_after=8)

add_page_break(doc)

# ── 8. gRPC ───────────────────────────────────────────────────────────────────
add_heading(doc, "8. Comunicación entre servicios — gRPC")

add_paragraph(doc, (
    "Con el objetivo de introducir conceptos de arquitectura de microservicios, se implementó "
    "un servicio independiente (notifications-service) que se comunica con el backend mediante "
    "el protocolo gRPC. Este patrón es habitual en sistemas de producción donde diferentes "
    "servicios necesitan comunicarse de forma eficiente y con contratos bien definidos."
), space_after=8)

add_heading(doc, "Contrato del servicio (Protocol Buffers)", level=2)

add_paragraph(doc, (
    "El contrato de la API gRPC se define en el fichero notifications.proto usando Protocol "
    "Buffers (protobuf), el lenguaje de serialización de Google. El servicio expone dos RPCs:"
), space_after=4)

add_bullet(doc, "Send(SendRequest) → SendResponse: envía una notificación a un destinatario por un canal específico (EMAIL, WHATSAPP, INTERNAL) y devuelve un identificador único y timestamp.")
add_bullet(doc, "GetStats(Empty) → StatsResponse: devuelve estadísticas de notificaciones enviadas, desglosadas por canal.")
add_paragraph(doc, "", space_after=4)

add_heading(doc, "Patrón API Gateway", level=2)

add_paragraph(doc, (
    "El backend actúa como API Gateway: expone endpoints REST hacia el frontend y, para las "
    "operaciones de notificaciones, traduce esas peticiones HTTP a llamadas gRPC hacia el "
    "microservicio. El frontend nunca accede directamente al microservicio gRPC, siempre a "
    "través del backend. Esta arquitectura tiene varias ventajas:"
), space_after=4)

add_bullet(doc, "El cliente (frontend) solo necesita conocer la API REST, no el protocolo gRPC.")
add_bullet(doc, "El backend puede aplicar autenticación y autorización antes de delegar al microservicio.")
add_bullet(doc, "El microservicio puede evolucionar independientemente sin afectar al frontend.")
add_paragraph(doc, "", space_after=8)

add_heading(doc, "Tecnologías gRPC", level=2)

add_paragraph(doc, (
    "Se utilizó la librería @grpc/grpc-js (implementación pure JavaScript de gRPC) junto con "
    "@grpc/proto-loader para cargar dinámicamente el fichero .proto en tiempo de ejecución. "
    "Esta aproximación evita el paso de compilación de protobuf y simplifica el flujo de "
    "desarrollo. Tanto el servidor (notifications-service) como el cliente (backend) cargan "
    "el mismo fichero .proto, garantizando que el contrato esté siempre sincronizado."
), space_after=8)

add_page_break(doc)

# ── 9. Git ────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Flujo de trabajo con Git y GitHub")

add_paragraph(doc, (
    "El proyecto sigue un flujo de trabajo basado en ramas que separa el desarrollo experimental "
    "del código estable."
), space_after=8)

add_heading(doc, "Estrategia de ramas", level=2)
add_bullet(doc, "main: rama de producción, siempre contiene código estable y funcional.")
add_bullet(doc, "develop: rama de integración donde se fusionan las nuevas funcionalidades antes de pasar a main.")
add_bullet(doc, "feature/*: ramas de funcionalidad específica (feature/login, feature/agenda, feature/observability-grpc, etc.).")
add_paragraph(doc, "", space_after=8)

add_heading(doc, "Convenciones de commits", level=2)
add_paragraph(doc, (
    "Se adoptó el estándar Conventional Commits para los mensajes de commit, usando prefijos "
    "semánticos que facilitan la generación automática de changelogs y comunican claramente "
    "el tipo de cambio:"
), space_after=4)

add_bullet(doc, "feat: nueva funcionalidad (feat: add agenda module CRUD)")
add_bullet(doc, "fix: corrección de error (fix: resolve JWT validation edge case)")
add_bullet(doc, "docs: cambios en documentación (docs: update README with Docker instructions)")
add_bullet(doc, "chore: tareas de mantenimiento (chore: add database seed)")
add_bullet(doc, "refactor: reestructuración sin cambio de comportamiento")
add_paragraph(doc, "", space_after=8)

add_heading(doc, "Pull Requests y revisión de código", level=2)
add_paragraph(doc, (
    "Cada funcionalidad se integra en develop mediante un Pull Request en GitHub, lo que "
    "permite al equipo revisar los cambios antes de fusionarlos. Este proceso sirve también "
    "como mecanismo de documentación: el historial de Pull Requests refleja la evolución del "
    "proyecto y las decisiones tomadas en cada fase."
), space_after=8)

add_heading(doc, "Acceso del profesor", level=2)
add_paragraph(doc, (
    "El profesor Javier del Valle Echevarri ha sido añadido como colaborador del repositorio "
    "GitHub (https://github.com/IzanUrrutikoetxea/GUDI-web-app), con acceso completo al "
    "código fuente, historial de commits, ramas y Pull Requests, así como a toda la "
    "documentación técnica del proyecto."
), space_after=8)

add_page_break(doc)

# ── 10. Conclusiones ──────────────────────────────────────────────────────────
add_heading(doc, "10. Conclusiones")

add_paragraph(doc, (
    "El desarrollo del proyecto GUDI ha permitido al equipo aplicar de forma práctica un "
    "conjunto amplio de conceptos y tecnologías del desarrollo web moderno. A lo largo de las "
    "seis fases del proyecto se ha construido una aplicación funcional y completa, desde la "
    "autenticación hasta la bandeja omnicanal, pasando por la gestión de citas y presupuestos."
), space_after=8)

add_paragraph(doc, (
    "La decisión de contenerizar toda la aplicación con Docker desde el inicio ha resultado "
    "especialmente valiosa, ya que ha permitido que cualquier miembro del equipo pueda levantar "
    "el entorno completo con un único comando, eliminando los clásicos problemas de "
    "incompatibilidades entre entornos de desarrollo."
), space_after=8)

add_paragraph(doc, (
    "La incorporación del stack de observabilidad ha sido una de las experiencias más "
    "enriquecedoras del proyecto. Implementar métricas con Prometheus y Grafana, trazas "
    "distribuidas con OpenTelemetry y Jaeger, y logs estructurados con el stack ELK ha "
    "proporcionado al equipo una comprensión práctica de cómo se monitoriza una aplicación "
    "real en producción. Como señaló el profesor, hasta que no se implementan estos sistemas "
    "de forma práctica resulta difícil entender completamente su funcionamiento y valor."
), space_after=8)

add_paragraph(doc, (
    "La implementación del microservicio de notificaciones con gRPC ha introducido conceptos "
    "de arquitectura de microservicios y comunicación entre servicios que complementan la "
    "formación en desarrollo web con una perspectiva más amplia de los sistemas distribuidos."
), space_after=8)

add_paragraph(doc, (
    "Como líneas de trabajo futuro, el proyecto podría enriquecerse con la implementación de "
    "una pipeline de CI/CD completa en GitHub Actions, la adición de tests automatizados "
    "(unitarios y de integración), la implementación de más microservicios comunicados por "
    "gRPC, o el despliegue en un entorno cloud como AWS o Azure utilizando Kubernetes para "
    "la orquestación de contenedores."
), space_after=8)

# ─── Save ─────────────────────────────────────────────────────────────────────
import os
output = os.path.join(os.path.dirname(__file__), "informe.docx")
doc.save(output)
print(f"Document saved to: {output}")
